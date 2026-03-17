import base64
import json
import os
import subprocess
from copy import deepcopy
from datetime import datetime
from typing import List, Dict, Any

import streamlit as st
from fpdf import FPDF  # kept for compatibility if needed later
from docx import Document


# BASE_DIR should point to the project root (folder that contains app.py),
# even though this file lives under features/resume/.
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
PROCESSED_DIR = os.path.join(BASE_DIR, "processed")
TEX_DIR = os.path.join(PROCESSED_DIR, "tex")
PDF_DIR = os.path.join(PROCESSED_DIR, "pdf")

os.makedirs(PROCESSED_DIR, exist_ok=True)
os.makedirs(TEX_DIR, exist_ok=True)
os.makedirs(PDF_DIR, exist_ok=True)

TEMPLATE_JSON_PATH = os.path.join(BASE_DIR, "resume_template.json")


def _load_json_config(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


CONFIG: Dict[str, Any] = _load_json_config(TEMPLATE_JSON_PATH)

LATEX_TEMPLATE_NAME: str = CONFIG.get("latex_template", "shamik_resume_template.tex")

DEFAULT_CONTACT: Dict[str, str] = CONFIG.get("contact", {})
TEMPLATES_BY_ROLE: Dict[str, Dict[str, Any]] = CONFIG.get("roles", {})
DEFAULT_EXPERIENCE: List[Dict[str, Any]] = CONFIG.get("experience", [])
DEFAULT_EDUCATION: List[Dict[str, str]] = CONFIG.get("education", [])
DEFAULT_SKILLS: List[str] = CONFIG.get("skills", [])
DEFAULT_LEADERSHIP: List[Dict[str, Any]] = CONFIG.get("leadership", [])

# Section order matches resume template: summary, education, experience, projects, leadership, skills
DEFAULT_SECTION_ORDER = ["summary", "education", "experience", "projects", "leadership", "skills"]
SECTION_TITLES = {
    "summary": "Professional Summary",
    "education": "Education",
    "experience": "Work Experience",
    "projects": "Projects",
    "leadership": "Leadership",
    "skills": "Technical Skills",
}


def get_default_word_template_path() -> str:
    return os.path.join(BASE_DIR, "default_resume_template.docx")


def load_docx_text_from_path(path: str) -> str:
    try:
        doc = Document(path)
    except Exception:
        return ""
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(p for p in paragraphs if p.strip())


def load_docx_text_from_file(file_obj) -> str:
    try:
        doc = Document(file_obj)
    except Exception:
        return ""
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(p for p in paragraphs if p.strip())


def _safe_for_multicell(text: str, max_chunk: int = 40) -> str:
    safe_words: list[str] = []
    for word in text.split(" "):
        if len(word) <= max_chunk:
            safe_words.append(word)
        else:
            chunks = [word[i : i + max_chunk] for i in range(0, len(word), max_chunk)]
            safe_words.append(" ".join(chunks))
    return " ".join(safe_words)


def latex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    for orig, repl in replacements.items():
        text = text.replace(orig, repl)
    return text


def extract_keywords_from_text(text: str) -> List[str]:
    tokens = [t.strip(".,:;()[]{}").lower() for t in text.split()]
    keep = [t for t in tokens if len(t) > 3]
    return sorted(list(set(keep)))


def score_project(project: Dict[str, Any], keywords: List[str]) -> int:
    score = 0
    text = (project.get("title", "") + " " + project.get("description", "")).lower()
    tags = " ".join(project.get("tags", [])).lower()
    full = text + " " + tags
    for kw in keywords:
        if kw.lower() in full:
            score += 1
    return score


def choose_projects(role: str, jd_text: str, extra_keywords: str, top_k: int = 3) -> List[Dict[str, Any]]:
    template = TEMPLATES_BY_ROLE.get(role, {})
    projects = template.get("projects", [])

    base_keywords = template.get("keywords", [])
    jd_keywords = extract_keywords_from_text(jd_text) if jd_text else []
    user_keywords = [k.strip().lower() for k in extra_keywords.split(",") if k.strip()] if extra_keywords else []

    all_keywords = list(set([*base_keywords, *jd_keywords, *user_keywords]))

    scored = [(score_project(p, all_keywords), p) for p in projects]
    scored.sort(key=lambda x: x[0], reverse=True)
    selected = [p for s, p in scored if s > 0]
    if len(selected) < top_k:
        remaining = [p for _, p in scored if p not in selected]
        selected.extend(remaining[: top_k - len(selected)])
    return selected[:top_k]


def build_summary(role: str, custom_summary: str | None = None) -> List[str]:
    template = TEMPLATES_BY_ROLE.get(role, {})
    default_summary = template.get("summary", [])
    if custom_summary:
        return [line.strip() for line in custom_summary.split("\n") if line.strip()]
    return default_summary


def default_file_name() -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"resume_{timestamp}.pdf"


def _needs_prefill_experience(items: Any) -> bool:
    if not isinstance(items, list):
        return True
    if not items and DEFAULT_EXPERIENCE:
        return True
    for it in items:
        if not isinstance(it, dict):
            return True
        if "company" not in it or "role" not in it or "period" not in it:
            return True
    return False


def _needs_prefill_education(items: Any) -> bool:
    if not isinstance(items, list):
        return True
    if not items and DEFAULT_EDUCATION:
        return True
    for it in items:
        if not isinstance(it, dict):
            return True
        if "degree" not in it or "institution" not in it or "year" not in it:
            return True
    return False


def _needs_prefill_leadership(items: Any) -> bool:
    if not isinstance(items, list):
        return True
    if not items and DEFAULT_LEADERSHIP:
        return True
    for it in items:
        if not isinstance(it, dict):
            return True
        if "title" not in it or "role" not in it or "period" not in it:
            return True
    return False


def _href_url(url: str) -> str:
    if not url:
        return ""
    return url.replace("%", r"\%").replace("#", r"\#")


def _build_header_tex(contact: Dict[str, str]) -> str:
    name = latex_escape((contact.get("name") or "").strip())
    if not name:
        name = "Your Name"
    loc = latex_escape((contact.get("location") or "").strip())
    phone = latex_escape((contact.get("phone") or "").strip())
    email = (contact.get("email") or "").strip()
    linkedin = (contact.get("linkedin") or "").strip()
    github = (contact.get("github") or "").strip()

    parts = []
    parts.append(r"\begin{center}")
    parts.append(rf"    {{\LARGE \textbf{{{name}}}}}\\[4pt]")
    line2_bits = [b for b in [loc, phone] if b]
    if email:
        mail_show = latex_escape(email)
        line2_bits.append(rf"\href{{mailto:{_href_url(email)}}}{{{mail_show}}}")
    if line2_bits:
        parts.append("    " + r" \textbar{} ".join(line2_bits) + r"\\")
    link_bits = []
    if linkedin:
        u = _href_url(linkedin)
        link_bits.append(rf"LinkedIn: \href{{{u}}}{{{latex_escape(linkedin)}}}")
    if github:
        u = _href_url(github)
        link_bits.append(rf"GitHub: \href{{{u}}}{{{latex_escape(github)}}}")
    if link_bits:
        parts.append("    " + r" \textbar{} ".join(link_bits))
    parts.append(r"\end{center}")
    return "\n".join(parts)


def _build_education_tex(education: List[Dict[str, Any]]) -> str:
    blocks: list[str] = []
    for edu in education:
        inst = latex_escape(edu.get("institution", ""))
        year = latex_escape(edu.get("year", ""))
        degree = latex_escape(edu.get("degree", ""))
        loc = latex_escape(edu.get("location", ""))
        if not inst and not degree:
            continue
        if inst and year:
            blocks.append(rf"\textbf{{{inst}}} \hfill {year}\\")
        elif inst:
            blocks.append(rf"\textbf{{{inst}}}\\")
        if degree or loc:
            if degree and loc:
                blocks.append(rf"{degree} \hfill {loc}\\[-2pt]")
            elif degree:
                blocks.append(degree + r"\\[-2pt]")
            else:
                blocks.append(loc + r"\\[-2pt]")
        bullets = edu.get("bullets") or edu.get("coursework")
        if isinstance(bullets, list) and bullets:
            blocks.append(r"\begin{itemize}")
            for b in bullets:
                blocks.append(rf"  \item {latex_escape(str(b))}")
            blocks.append(r"\end{itemize}")
        blocks.append("")
    return "\n".join(blocks).strip() or "% No education entries"


def _build_experience_tex(experience: List[Dict[str, Any]]) -> str:
    blocks: list[str] = []
    for exp in experience:
        role = latex_escape(exp.get("role", ""))
        company = latex_escape(exp.get("company", ""))
        period = latex_escape(exp.get("period", ""))
        loc = latex_escape(exp.get("location", ""))
        header = rf"\textbf{{{role}}} \textbar{{}} {company}"
        if loc:
            header += rf" \textbar{{}} {loc}"
        if period:
            header += rf" \hfill {period}"
        blocks.append(header + r"\\[-2pt]")
        blocks.append(r"\begin{itemize}")
        for h in exp.get("highlights") or []:
            blocks.append(rf"  \item {latex_escape(str(h))}")
        blocks.append(r"\end{itemize}")
        blocks.append("")
    return "\n".join(blocks).strip() or "% No experience entries"


def _build_leadership_tex(leadership: List[Dict[str, Any]]) -> str:
    if not leadership:
        return "% No leadership entries"
    blocks: list[str] = []
    for item in leadership:
        title = latex_escape(item.get("title", ""))
        role = latex_escape(item.get("role", ""))
        period = latex_escape(item.get("period", ""))
        line = rf"\textbf{{{title}}}"
        if role:
            line += rf" \textbar{{}} {role}"
        if period:
            line += rf" \hfill {period}"
        blocks.append(line + r"\\[-2pt]")
        blocks.append(r"\begin{itemize}")
        for b in item.get("bullets") or []:
            blocks.append(rf"  \item {latex_escape(str(b))}")
        blocks.append(r"\end{itemize}")
        blocks.append("")
    return "\n".join(blocks).strip()


def _build_skills_tex(skills: List[str]) -> str:
    if not skills:
        return "% No skills listed"
    line = ", ".join(latex_escape(s) for s in skills if s)
    return rf"\textbf{{Skills:}} {line}"


def _build_projects_tex(projects: List[Dict[str, Any]]) -> str:
    project_blocks: list[str] = []
    for proj in projects:
        title = latex_escape(proj.get("title", ""))
        subtitle = latex_escape(proj.get("subtitle", ""))
        bullets = proj.get("bullets") or []
        lines: list[str] = []
        if subtitle:
            lines.append(rf"\textbf{{{title}}} \hfill {subtitle}\\[-2pt]")
        else:
            lines.append(rf"\textbf{{{title}}}\\[-2pt]")
        lines.append(r"\begin{itemize}")
        if bullets:
            for b in bullets:
                lines.append(rf"  \item {latex_escape(b)}")
        else:
            tech = proj.get("tech", "")
            desc = proj.get("description", "")
            if tech:
                lines.append(rf"  \item \textit{{{latex_escape(tech)}}}")
            if desc:
                lines.append(rf"  \item {latex_escape(desc)}")
        lines.append(r"\end{itemize}")
        project_blocks.append("\n".join(lines))
    return "\n\n".join(project_blocks) if project_blocks else "% No projects selected"


def generate_pdf(
    contact: Dict[str, str],
    summary_lines: List[str],
    experience: List[Dict[str, Any]],
    projects: List[Dict[str, Any]],
    education: List[Dict[str, Any]],
    skills: List[str],
    target_path: str,
    leadership: List[Dict[str, Any]] | None = None,
    base_text: str | None = None,
    use_word_template: bool = False,
    section_order: List[str] | None = None,
) -> None:
    template_path = os.path.join(BASE_DIR, LATEX_TEMPLATE_NAME)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"LaTeX template not found at {template_path}")

    with open(template_path, "r", encoding="utf-8") as f:
        tex_template = f.read()

    order = section_order or DEFAULT_SECTION_ORDER

    header_tex = _build_header_tex(contact)

    summary_text = " ".join(line.strip() for line in summary_lines if line.strip())
    summary_tex = latex_escape(summary_text) if summary_text else "% No summary"

    education_tex = _build_education_tex(education)
    experience_tex = _build_experience_tex(experience)
    leadership_tex = _build_leadership_tex(leadership or DEFAULT_LEADERSHIP)
    skills_tex = _build_skills_tex(skills)
    projects_tex = _build_projects_tex(projects)

    section_tex_map = {
        "summary": summary_tex,
        "education": education_tex,
        "experience": experience_tex,
        "projects": projects_tex,
        "leadership": leadership_tex,
        "skills": skills_tex,
    }

    body_parts = []
    for sec_id in order:
        if sec_id in section_tex_map:
            title = SECTION_TITLES.get(sec_id, sec_id.title())
            body_parts.append(f"\\section*{{{title}}}\n{section_tex_map[sec_id]}")

    body_tex = "\n\n".join(body_parts)

    filled_tex = tex_template.replace("%%HEADER%%", header_tex).replace("%%BODY%%", body_tex)

    base_name = os.path.splitext(os.path.basename(target_path))[0]
    tex_out_path = os.path.join(TEX_DIR, f"{base_name}.tex")
    with open(tex_out_path, "w", encoding="utf-8") as f:
        f.write(filled_tex)

    try:
        subprocess.run(
            [
                "xelatex",
                "-interaction=nonstopmode",
                "-halt-on-error",
                f"-jobname={base_name}",
                os.path.basename(tex_out_path),
            ],
            cwd=TEX_DIR,
            check=True,
            capture_output=True,
            text=True,
        )
    except FileNotFoundError as e:
        raise RuntimeError(
            "LaTeX engine 'xelatex' was not found. "
            "Please install a TeX distribution (e.g. MiKTeX or TeX Live) and ensure 'xelatex' is on your PATH."
        ) from e
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"LaTeX compilation failed: {e.stderr}") from e

    compiled_pdf = os.path.join(TEX_DIR, f"{base_name}.pdf")
    if os.path.exists(compiled_pdf):
        os.replace(compiled_pdf, target_path)

    for ext in (".tex", ".aux", ".log", ".out"):
        aux_path = os.path.join(TEX_DIR, f"{base_name}{ext}")
        try:
            if os.path.exists(aux_path):
                os.remove(aux_path)
        except OSError:
            pass


def run() -> None:
    """Render the resume builder UI."""
    st.title("Custom Resume Builder")

    st.markdown(
        """
This tool helps you quickly generate a tailored resume PDF from a default profile.

- Choose the **target role / template**
- Paste a **job description** and/or provide **keywords**
- Adjust summary and selected projects if needed
- Export as a **PDF** saved under `processed/pdf/resume_YYYYMMDD_HHMMSS.pdf`
"""
    )

    with st.sidebar:
        st.header("Target Role & Inputs")
        role = st.selectbox("Target role / template", list(TEMPLATES_BY_ROLE.keys()))

        jd_source = st.radio(
            "Job description input type",
            ["Text box", "Upload .txt file"],
            index=0,
        )

        jd_text = ""
        if jd_source == "Text box":
            jd_text = st.text_area(
                "Paste job description (optional but recommended)",
                height=200,
            )
        else:
            uploaded_file = st.file_uploader("Upload job description (.txt)", type=["txt"])
            if uploaded_file is not None:
                jd_text = uploaded_file.read().decode("utf-8", errors="ignore")

        extra_keywords = st.text_input(
            "Extra keywords (comma separated)",
            help="e.g. NLP, recommendation systems, time series",
        )

        auto_projects = st.checkbox(
            "Auto-select relevant projects from template",
            value=True,
            help="If unchecked, you can manually pick projects.",
        )

    if "experience_list" not in st.session_state or _needs_prefill_experience(st.session_state.experience_list):
        st.session_state.experience_list = deepcopy(DEFAULT_EXPERIENCE)
    if "education_list" not in st.session_state or _needs_prefill_education(st.session_state.education_list):
        st.session_state.education_list = deepcopy(DEFAULT_EDUCATION)
    if "leadership_list" not in st.session_state or _needs_prefill_leadership(st.session_state.leadership_list):
        st.session_state.leadership_list = deepcopy(DEFAULT_LEADERSHIP)
    if "section_order" not in st.session_state:
        st.session_state.section_order = DEFAULT_SECTION_ORDER.copy()
    if "skills_text" not in st.session_state:
        st.session_state.skills_text = ", ".join(DEFAULT_SKILLS)

    experience_list = st.session_state.experience_list
    education_list = st.session_state.education_list
    leadership_list = st.session_state.leadership_list
    section_order = st.session_state.section_order
    skills_text_state = st.session_state.skills_text

    col_main, col_preview = st.columns([1, 1])

    with col_main:
        st.subheader("Edit resume (all sections drive the PDF)")

        contact = deepcopy(DEFAULT_CONTACT)
        with st.expander("Contact details", expanded=True):
            contact["name"] = st.text_input("Name", contact.get("name", ""), key="contact_name")
            contact["email"] = st.text_input("Email", contact.get("email", ""), key="contact_email")
            contact["phone"] = st.text_input("Phone", contact.get("phone", ""), key="contact_phone")
            contact["location"] = st.text_input("Location", contact.get("location", ""), key="contact_location")
            contact["linkedin"] = st.text_input("LinkedIn URL", contact.get("linkedin", ""), key="contact_linkedin")
            contact["github"] = st.text_input("GitHub URL", contact.get("github", ""), key="contact_github")

        st.markdown("---")

        st.subheader("Section order")
        st.caption("Use ↑ ↓ to change the order of sections in your resume. Changes reflect in preview.")
        for i, sec_id in enumerate(section_order):
            sec_title = SECTION_TITLES.get(sec_id, sec_id.title())
            c1, c2, c3 = st.columns([4, 1, 1])
            with c1:
                st.markdown(f"**{i + 1}. {sec_title}**")
            with c2:
                if st.button("↑", key=f"sec_up_{sec_id}", disabled=(i == 0)):
                    if i > 0:
                        section_order[i], section_order[i - 1] = section_order[i - 1], section_order[i]
                        st.session_state.section_order = section_order
                        st.rerun()
            with c3:
                if st.button("↓", key=f"sec_down_{sec_id}", disabled=(i == len(section_order) - 1)):
                    if i < len(section_order) - 1:
                        section_order[i], section_order[i + 1] = section_order[i + 1], section_order[i]
                        st.session_state.section_order = section_order
                        st.rerun()

        st.markdown("---")

        default_summary_text = "\n".join(build_summary(role, None))

        with st.expander("1. Professional Summary", expanded=True):
            summary_override = st.text_area(
                "Summary",
                default_summary_text,
                height=100,
                key=f"summary_override_{role}",
                help="Override the role-based summary. Left empty uses the template summary.",
            )
            summary_lines = build_summary(role, summary_override if summary_override.strip() else None)

        with st.expander("2. Education", expanded=True):
            for ei, edu in enumerate(education_list):
                with st.container():
                    c1, c2 = st.columns([5, 1])
                    with c1:
                        st.markdown("**Entry " + str(ei + 1) + "**")
                    with c2:
                        if st.button("Remove", key=f"edu_remove_{ei}", type="secondary"):
                            st.session_state.education_list.pop(ei)
                            st.rerun()
                    edu["degree"] = st.text_input("Degree", edu.get("degree", ""), key=f"edu_deg_{ei}")
                    edu["institution"] = st.text_input("Institution", edu.get("institution", ""), key=f"edu_inst_{ei}")
                    edu["year"] = st.text_input("Year", edu.get("year", ""), key=f"edu_year_{ei}")
                    edu["location"] = st.text_input("Location", edu.get("location", ""), key=f"edu_loc_{ei}")
            if st.button("+ Add education", key="edu_add"):
                st.session_state.education_list.append(
                    {
                        "degree": "",
                        "institution": "",
                        "year": "",
                        "location": "",
                    }
                )
                st.rerun()

        with st.expander("3. Work Experience", expanded=True):
            for i, exp in enumerate(experience_list):
                with st.container():
                    c1, c2 = st.columns([5, 1])
                    with c1:
                        st.markdown("**Entry " + str(i + 1) + "**")
                    with c2:
                        if st.button("Remove", key=f"exp_remove_{i}", type="secondary"):
                            st.session_state.experience_list.pop(i)
                            st.rerun()
                    exp["company"] = st.text_input("Company", exp.get("company", ""), key=f"exp_company_{i}")
                    exp["role"] = st.text_input("Role", exp.get("role", ""), key=f"exp_role_{i}")
                    exp["period"] = st.text_input("Period", exp.get("period", ""), key=f"exp_period_{i}")
                    new_highlights = []
                    for j, h in enumerate(list(exp.get("highlights") or [])):
                        new_h = st.text_input(f"Bullet {j + 1}", h, key=f"exp_{i}_h_{j}")
                        if new_h.strip():
                            new_highlights.append(new_h.strip())
                    if st.button("+ Add bullet", key=f"exp_add_bullet_{i}"):
                        hl = exp.get("highlights") or []
                        hl.append("")
                        exp["highlights"] = hl
                        st.rerun()
                    exp["highlights"] = new_highlights
            if st.button("+ Add experience", key="exp_add"):
                st.session_state.experience_list.append(
                    {
                        "company": "",
                        "role": "",
                        "period": "",
                        "highlights": [],
                    }
                )
                st.rerun()

        with st.expander("4. Leadership", expanded=True):
            for li, item in enumerate(leadership_list):
                with st.container():
                    c1, c2 = st.columns([5, 1])
                    with c1:
                        st.markdown("**Entry " + str(li + 1) + "**")
                    with c2:
                        if st.button("Remove", key=f"lead_remove_{li}", type="secondary"):
                            st.session_state.leadership_list.pop(li)
                            st.rerun()
                    item["title"] = st.text_input("Title / Organization", item.get("title", ""), key=f"lead_title_{li}")
                    item["role"] = st.text_input("Role", item.get("role", ""), key=f"lead_role_{li}")
                    item["period"] = st.text_input("Period", item.get("period", ""), key=f"lead_period_{li}")
                    new_bullets = []
                    for j, b in enumerate(list(item.get("bullets") or [])):
                        new_b = st.text_input(f"Bullet {j + 1}", b, key=f"lead_{li}_b_{j}")
                        if new_b.strip():
                            new_bullets.append(new_b.strip())
                    if st.button("+ Add bullet", key=f"lead_add_bullet_{li}"):
                        bl = item.get("bullets") or []
                        bl.append("")
                        item["bullets"] = bl
                        st.rerun()
                    item["bullets"] = new_bullets
            if st.button("+ Add leadership", key="lead_add"):
                st.session_state.leadership_list.append(
                    {
                        "title": "",
                        "role": "",
                        "period": "",
                        "bullets": [],
                    }
                )
                st.rerun()

        with st.expander("5. Technical Skills", expanded=False):
            skills_text = st.text_area(
                "Skills (comma separated)",
                skills_text_state,
                key="skills_area",
            )
            st.session_state.skills_text = skills_text
            skills = [s.strip() for s in skills_text.split(",") if s.strip()]

    with col_preview:
        st.subheader("Projects for This Resume")
        template_projects = TEMPLATES_BY_ROLE[role]["projects"]

        if auto_projects:
            selected_projects = choose_projects(role, jd_text, extra_keywords, top_k=3)
            st.markdown("Auto-selected based on job description & keywords:")
            for p in selected_projects:
                desc = ""
                if "description" in p:
                    desc = p["description"]
                elif "bullets" in p and p["bullets"]:
                    desc = p["bullets"][0]
                st.markdown(f"- **{p.get('title', '')}**  \n  {desc}")
        else:
            selected_indices = []
            for i, proj in enumerate(template_projects):
                checked = st.checkbox(
                    f"{proj['title']}",
                    value=True,
                    help=proj["description"],
                    key=f"proj_{role}_{i}",
                )
                if checked:
                    selected_indices.append(i)
            selected_projects = [template_projects[i] for i in selected_indices]

        st.markdown("---")
        st.subheader("Generate / Preview PDF")

        preview_path = None
        if selected_projects:
            preview_name = "preview_resume.pdf"
            preview_path = os.path.join(PDF_DIR, preview_name)
            generate_pdf(
                contact=contact,
                summary_lines=summary_lines,
                experience=experience_list,
                projects=selected_projects,
                education=education_list,
                skills=skills,
                leadership=leadership_list,
                target_path=preview_path,
                base_text=None,
                use_word_template=False,
                section_order=section_order,
            )

            try:
                with open(preview_path, "rb") as f:
                    base64_pdf = base64.b64encode(f.read()).decode("utf-8")
                pdf_display = (
                    f'<iframe src="data:application/pdf;base64,{base64_pdf}" '
                    'width="100%" height="95vh" style="min-height:900px;border:none;" '
                    'type="application/pdf"></iframe>'
                )
                st.markdown("**Live PDF preview** — updates on every change:", unsafe_allow_html=False)
                st.markdown(pdf_display, unsafe_allow_html=True)
            except OSError:
                st.warning("Preview PDF could not be loaded. Please try adjusting a field again.")
        else:
            st.info("Select at least one project to see a live PDF preview.")

        st.markdown("---")

        if st.button("Build final tailored resume PDF"):
            if not selected_projects:
                st.error("Please ensure at least one project is selected.")
            else:
                file_name = default_file_name()
                target_path = os.path.join(PDF_DIR, file_name)

                generate_pdf(
                    contact=contact,
                    summary_lines=summary_lines,
                    experience=experience_list,
                    projects=selected_projects,
                    education=education_list,
                    skills=skills,
                    leadership=leadership_list,
                    target_path=target_path,
                    base_text=None,
                    use_word_template=False,
                    section_order=section_order,
                )

                st.success(f"Final PDF generated and saved as `{file_name}` in `processed/pdf/` folder.")
                with open(target_path, "rb") as f:
                    st.download_button(
                        "Download final PDF",
                        data=f,
                        file_name=file_name,
                        mime="application/pdf",
                    )

